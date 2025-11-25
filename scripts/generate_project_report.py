from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def set_page_to_a4(document: Document):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)


def configure_styles(document: Document):
    normal_style = document.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)

    h1 = document.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True

    h2 = document.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_bullets(document: Document, items):
    for item in items:
        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def add_image_placeholder(document: Document, caption: str = None):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Insert Image Here]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if caption:
        cap = document.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report(document: Document):
    # Title
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('AgriSmart – AI-Powered Crop Yield Prediction Platform')
    title_run.bold = True
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)

    # Overview
    add_heading(document, 'Project Overview', level=1)
    add_paragraph(document, (
        'AgriSmart is a comprehensive agricultural intelligence platform leveraging AI/ML to deliver '
        'data-driven insights for optimal crop management. It combines real-time weather data, soil '
        'analysis, pest detection, and market analytics to maximize yield and profitability.'
    ))

    # Core Features
    add_heading(document, 'Key Features', level=1)
    add_paragraph(document, 'Core Modules:', bold=True)
    add_bullets(document, [
        'AI-Powered Crop Yield Prediction',
        'Profitable Crops Predictor',
        'Soil Health Monitor (NPK, pH, health scoring)',
        'Pest & Disease Detection (Computer Vision)',
        'Weather Intelligence & Forecasting',
        'Smart Irrigation Management',
        'Advanced Analytics & Reporting',
        'Multi-language Support (English, Hindi, Gujarati)'
    ])

    add_paragraph(document, 'AI/ML Capabilities:', bold=True)
    add_bullets(document, [
        'Custom ML models for yield forecasting and recommendations',
        'Deep learning-based pest classification',
        'Predictive analytics for weather patterns',
        'Market intelligence and profit optimization',
        'Risk assessment including weather and market volatility'
    ])

    # Architecture
    add_heading(document, 'Technical Architecture', level=1)
    add_paragraph(document, 'Frontend (React 18, Vite, TailwindCSS, Redux Toolkit, React Router v6, i18next, Framer Motion)')
    add_paragraph(document, 'Backend (FastAPI, Uvicorn, Pydantic, JWT Authentication, CORS)')
    add_paragraph(document, 'AI/ML (Scikit-learn, TensorFlow/Keras, NumPy, Pandas, OpenCV, Joblib)')
    add_paragraph(document, 'External APIs (OpenWeatherMap, Geolocation services, Market data APIs)')

    # Implementation Highlights (Recent Work)
    add_heading(document, 'Implementation Highlights', level=1)
    add_paragraph(document, 'Authentication Enhancements:', bold=True)
    add_bullets(document, [
        'Added Account Recovery and New Password pages in frontend',
        'Integrated service methods for recovery and password update',
        'Implemented API endpoints for account recovery and credential update',
        'Secure tokens via JWT; server-side validation and hashing'
    ])

    # Usage & Setup
    add_heading(document, 'Setup and Usage', level=1)
    add_paragraph(document, 'Prerequisites:', bold=True)
    add_bullets(document, [
        'Node.js v16+',
        'Python 3.8+',
        'npm or yarn',
        'Git'
    ])
    add_paragraph(document, 'Frontend:', bold=True)
    add_bullets(document, [
        'Install: npm install',
        'Run dev server: npm run dev (http://localhost:3001)'
    ])
    add_paragraph(document, 'Backend:', bold=True)
    add_bullets(document, [
        'Install: pip install -r requirements.txt',
        'Run API: python app/main.py (http://localhost:8000)'
    ])
    add_paragraph(document, 'Environment Configuration (.env):', bold=True)
    add_bullets(document, [
        'REACT_APP_WEATHER_API_KEY=your_openweather_api_key',
        'VITE_BACKEND_URL=http://localhost:8000',
        'JWT_SECRET_KEY=your_jwt_secret_key'
    ])

    # Screenshots Placeholders
    add_heading(document, 'Screenshots', level=1)
    add_paragraph(document, 'Dashboard Overview:', bold=True)
    add_image_placeholder(document, caption='Dashboard – Yield Prediction & Analytics')
    add_paragraph(document, 'Account Recovery Flow:', bold=True)
    add_image_placeholder(document, caption='Account Recovery – Email Submission')
    add_paragraph(document, 'New Password Flow:', bold=True)
    add_image_placeholder(document, caption='New Password – Token Validation & Update')

    # API Examples
    add_heading(document, 'API Examples', level=1)
    add_paragraph(document, 'Crop Recommendation (POST /api/crop-prediction/recommend):', bold=True)
    add_paragraph(document, (
        '{ "nitrogen": 90.0, "phosphorus": 42.0, "potassium": 43.0, "temperature": 20.87, '
        '"humidity": 82.0, "ph": 6.5, "rainfall": 202.93 }'
    ))
    add_paragraph(document, 'Authentication:', bold=True)
    add_bullets(document, [
        'POST /api/auth/register',
        'POST /api/auth/login',
        'POST /api/auth/forgot-password',
        'POST /api/auth/reset-password'
    ])

    # Technology Stack Summary
    add_heading(document, 'Technology Stack Summary', level=1)
    add_bullets(document, [
        'Frontend: React 18, Vite, TailwindCSS, Redux Toolkit, React Router v6',
        'Backend: FastAPI, Uvicorn, Pydantic, JWT, CORS',
        'AI/ML: Scikit-learn, TensorFlow/Keras, NumPy, Pandas, OpenCV',
        'Visualization: D3.js, Recharts',
        'Icons & UI: Lucide React, Framer Motion',
        'Internationalization: i18next'
    ])

    # Contribution & License
    add_heading(document, 'Contribution & License', level=1)
    add_paragraph(document, 'Contribution Guidelines:', bold=True)
    add_bullets(document, [
        'Fork repository and create feature branch',
        'Follow code style and add tests where applicable',
        'Open a detailed pull request'
    ])
    add_paragraph(document, 'License:', bold=True)
    add_paragraph(document, 'MIT License – see repository LICENSE file for details.')

    # Conclusion
    add_heading(document, 'Conclusion', level=1)
    add_paragraph(document, (
        'AgriSmart delivers actionable, AI-driven insights to enhance farm productivity and sustainability. '
        'This report summarizes the architecture, features, setup, and recent authentication enhancements. '
        'Screenshots can be added where placeholders are indicated.'
    ))


def main():
    document = Document()
    set_page_to_a4(document)
    configure_styles(document)
    build_report(document)

    # Save to project root
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    output_path = os.path.join(root, 'AgriSmart_Project_Report.docx')
    try:
        document.save(output_path)
        print(f'Report generated: {output_path}')
    except PermissionError:
        # If file is open/locked, save to an alternate filename
        alt_output = os.path.join(root, 'AgriSmart_Project_Report_updated.docx')
        document.save(alt_output)
        print(f'Report generated (alternate file due to lock): {alt_output}')


if __name__ == '__main__':
    main()